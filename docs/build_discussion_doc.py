from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "博弈框架与数据集构建讨论稿.md"
OUTPUT = ROOT / "博弈框架与数据集构建讨论稿.docx"

BLUE = "215A8E"
DARK = "17324D"
LIGHT = "EAF2F8"
MID = "5E6E7E"
INK = "20262E"


def set_font(run, name="Arial Unicode MS", size=11, bold=None, color=INK, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill, border=None):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    if border:
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:color"), border)
        left.set(qn("w:space"), "8")
        pbdr.append(left)
        ppr.append(pbdr)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MID)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    endrun = paragraph.add_run(" 页")
    set_font(endrun, size=9, color=MID)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, DARK, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("研究与工程讨论稿")
    set_font(r, size=11, bold=True, color=BLUE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Hyper-Decept 博弈框架\n与数据集构建方案")
    set_font(r, size=28, bold=True, color=DARK)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("从开放环仿真到可追溯、解释引导的闭环攻防博弈")
    set_font(r, size=13, color=MID)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    widths = [Inches(1.45), Inches(4.55)]
    rows = [
        ("文档状态", "v0.1 · 内部讨论"),
        ("形成日期", "2026-07-23"),
        ("目标截稿", "2026-09-10"),
        ("讨论目标", "冻结创新边界、主实验矩阵、结果留痕与近期分工"),
    ]
    for row, values in zip(table.rows, rows):
        for i, (cell, value) in enumerate(zip(row.cells, values)):
            cell.width = widths[i]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=10.5, bold=(i == 0), color=DARK if i == 0 else INK)
            if i == 0:
                cell._tc.get_or_add_tcPr().append(
                    OxmlElement("w:shd")
                )
                cell._tc.tcPr[-1].set(qn("w:fill"), LIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("内部材料 · 结论需由可复现实验与文献检索支持")
    set_font(r, size=9.5, italic=True, color=MID)
    doc.add_page_break()


def add_rich_text(paragraph, text, base_size=10.5):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Courier New", size=9.2, color=DARK)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=base_size)


def add_body(doc, lines):
    in_quote = False
    for raw in lines:
        line = raw.rstrip()
        if not line or line.startswith("版本：") or line.startswith("日期：") or line.startswith("目标截稿："):
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            title = line[3:]
            if title == "一页结论":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run(title)
                set_font(r, size=17, bold=True, color=DARK)
            else:
                doc.add_paragraph(title, style="Heading 1")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            shade_paragraph(p, LIGHT, BLUE)
            add_rich_text(p, line[2:], 11)
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\.\s+", "", line))
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
            continue
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(4)
            add_rich_text(p, line, 11)
            continue
        p = doc.add_paragraph()
        add_rich_text(p, line)


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("HYPER-DECEPT · INTERNAL DISCUSSION")
    set_font(r, size=8.5, bold=True, color=MID)
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)
    add_body(doc, lines)

    props = doc.core_properties
    props.title = "Hyper-Decept 博弈框架与数据集构建方案"
    props.subject = "闭环攻防博弈、仿真数据集、TwiBot-22 与实验留痕"
    props.author = "Hyper-Decept 项目组"
    props.keywords = "Hyper-Decept, multi-agent game, dataset, explainability, TwiBot-22"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
